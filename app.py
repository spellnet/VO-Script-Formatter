"""
PP Script Timecoder — Web App
Flask backend: upload script + video → timecoded .docx via OpenAI Whisper API
"""

import os, re, uuid, json, time, shutil, tempfile, threading, subprocess, csv
from datetime import datetime, timezone
from pathlib import Path
from difflib import SequenceMatcher

from flask import (Flask, request, jsonify, send_file,
                   render_template, Response)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024 * 1024  # 2 GB upload limit

# ── Job store (in-memory — fine for single-server use) ────────────────────────
jobs = {}   # job_id → {"status", "progress", "log", "output_path", "error"}


# ─────────────────────────────────────────────────────────────────────────────
# Utility helpers
# ─────────────────────────────────────────────────────────────────────────────

def normalize(text):
    text = text.lower()
    text = re.sub(r"[^\w\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()

def similarity(a, b):
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()

def seconds_to_tc(secs, offset_secs=0.0, fps=25):
    """HH:MM:SS:FF — full timecode with frames, floored so TC is never late."""
    total = max(0.0, secs + offset_secs)
    h  = int(total // 3600)
    m  = int((total % 3600) // 60)
    s  = int(total % 60)
    ff = int((total - int(total)) * fps)   # floor frames, never late
    if ff >= fps:
        ff = fps - 1
    return f"{h:02d}:{m:02d}:{s:02d}:{ff:02d}"

def tc_to_seconds(tc):
    parts = tc.strip().replace(";", ":").split(":")
    try:
        if len(parts) == 4:
            return int(parts[0])*3600 + int(parts[1])*60 + int(parts[2]) + int(parts[3])/25.0
        elif len(parts) == 3:
            return int(parts[0])*3600 + int(parts[1])*60 + int(parts[2])
    except:
        pass
    return 0.0

def dur_str(secs):
    """HH:MM:SS:FF — full duration with frames, floored."""
    if secs <= 0:
        return ""
    h  = int(secs // 3600)
    m  = int((secs % 3600) // 60)
    s  = int(secs % 60)
    ff = int((secs - int(secs)) * 25)
    if ff >= 25:
        ff = 24
    return f"{h:02d}:{m:02d}:{s:02d}:{ff:02d}"


# ─────────────────────────────────────────────────────────────────────────────
# Script parser  (same logic as the desktop app)
# ─────────────────────────────────────────────────────────────────────────────

def parse_source_script(docx_path):
    """
    Parse a two-column Stampede source script:
      Col 0 (left)  — scene name OR editorial note
      Col 1 (right) — VO (bold uppercase) or Actuality (italic / SPEAKER:)

    Rules:
    - First table row = document title block → SKIP entirely
    - Col 0 bold/uppercase/short → section header row (no TC, blue)
    - Col 0 other text → editorial note, attach to Notes of next content row
    - Col 1 bold uppercase → VO line
    - Col 1 italic or SPEAKER: → Actuality line
    - Section/part labels that appear IN col 1 (e.g. "COMM", "PART TWO")
      → section header row, NOT a VO line with a TC
    """
    from docx import Document

    p = Path(docx_path)
    if not p.exists():
        raise RuntimeError(f"File not found: {docx_path}")
    if p.stat().st_size < 4096:
        raise RuntimeError(
            f"File appears to be a OneDrive placeholder ({p.stat().st_size} bytes). "
            "Right-click → 'Always keep on this device', then retry."
        )
    import tempfile, shutil
    tmp = Path(tempfile.mkdtemp()) / "source.docx"
    shutil.copy2(str(p), str(tmp))
    try:
        doc = Document(str(tmp))
    except Exception as e:
        raise RuntimeError(f"Cannot open script (is it open in Word?): {e}")

    lines = []

    def is_bold(para):
        # Check run-level bold
        run_bold = any(r.bold for r in para.runs if r.text.strip())
        if run_bold:
            return True
        # Check paragraph style bold
        if para.style and para.style.font and para.style.font.bold:
            return True
        return False

    def is_italic(para):
        return any(r.italic for r in para.runs if r.text.strip())

    # Words that mark a label row rather than VO content
    LABEL_KEYWORDS = {
        "PART ONE","PART TWO","PART THREE","PART FOUR","PART FIVE",
        "PART SIX","PART BREAK","ACT 1","ACT 2","ACT 3","ACT 4","ACT 5",
        "ACT 6","TITLES","COLD OPEN","TEASER","TAG","END OF SHOW",
        "WEIGHT LOSS","CODA","COMM","RECAP",
    }

    def is_label(text):
        t = text.upper().strip()
        if any(kw in t for kw in LABEL_KEYWORDS):
            return True
        # Short bold all-caps (≤6 words) with no lowercase = label
        if t == t.upper() and len(text.split()) <= 6 and len(text) > 2:
            return True
        return False

    def classify_content_cell(cell, pending_note=""):
        """
        The single rule: BOLD = VO. NOT BOLD = Actuality.

        Exceptions (skip entirely, produce no output row):
          - COMM or COMM:  (script formatting marker)
          - Bare "VO:" or "VO"  (label before VO text — next bold para IS the VO)
          - Bare "Actuality:"  (label before actuality text)
          - Empty paragraphs

        Speaker labels (OFFICER:, WOMAN: etc.) are not bold → actuality.
        "Sync" and other non-bold cues → actuality.
        """
        import re as _re
        SPEAKER_RE = _re.compile(r'^([A-Z][A-Z0-9 .\\-]{0,30}):[\s]*(.*)', _re.IGNORECASE)

        results  = []
        paras    = cell.paragraphs
        n        = len(paras)
        i        = 0

        while i < n:
            para = paras[i]
            text = para.text.strip()
            if not text:
                i += 1; continue

            bold   = is_bold(para)
            italic = is_italic(para)

            # ── Skip markers ──────────────────────────────────────────────
            t_upper = text.upper().rstrip(':').strip()
            if t_upper == 'COMM':
                i += 1; continue
            if t_upper == 'VO' and bold:
                i += 1; continue   # bare VO: label — next para is the real VO
            if text.lower().rstrip(':').strip() == 'actuality':
                i += 1; continue   # bare Actuality: label

            # ── BOLD = VO ─────────────────────────────────────────────────
            if bold:
                # Each bold paragraph is its own VO line (separate EDL event)
                results.append({
                    'type':    'vo',
                    'speaker': None,
                    'text':    text,
                    'notes':   '',
                })
                i += 1; continue

            # ── NOT BOLD = Actuality ──────────────────────────────────────
            # Check for SPEAKER: pattern
            m = SPEAKER_RE.match(text)
            if m and m.group(1).upper() == m.group(1):
                # e.g. "OFFICER: Hello" or bare "OFFICER:" followed by dialogue
                speaker = m.group(1).strip()
                diag    = m.group(2).strip()
                block   = [diag] if diag else []
                # Collect following non-bold lines as this speaker's dialogue
                while i + 1 < n:
                    nxt = paras[i + 1]
                    nt  = nxt.text.strip()
                    if not nt:
                        i += 1; continue
                    if is_bold(nxt):   # next bold line = new VO
                        break
                    nm = SPEAKER_RE.match(nt)
                    if nm and nm.group(1).upper() == nm.group(1):
                        break          # next speaker label
                    block.append(nt); i += 1
                results.append({
                    'type':    'act',
                    'speaker': speaker,
                    'text':    ' '.join(block),
                    'notes':   '',
                })
            else:
                # Plain actuality / sync / stage direction
                results.append({
                    'type':    'act',
                    'speaker': None,
                    'text':    text,
                    'notes':   '',
                })
            i += 1

        for r in results:
            r.setdefault('notes', '')
        return results

    # ── Find script table ─────────────────────────────────────────────────────
    script_table = None
    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(
            c.text.strip().lower() for c in table.rows[0].cells
        )
        TC_KEYWORDS = {"time code", "timecode", "tc in", "tc out"}
        if any(kw in first_row_text for kw in TC_KEYWORDS):
            continue
        if script_table is None or len(table.rows) > len(script_table.rows):
            script_table = table

    if script_table is None:
        raise RuntimeError("Could not find a script table in the document.")

    # ── Walk rows ─────────────────────────────────────────────────────────────
    current_section = None
    first_row       = True

    for row in script_table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue

        left  = cells[0].text.strip().strip("*").strip()
        right = cells[1].text.strip()

        # Always skip the first table row — it is always the title block
        # (programme name, version, date, PL SCRIPT etc.)
        if first_row:
            first_row = False
            continue

        # Left column — ONLY use for scene/segment header names.
        # ALL editorial notes, SG comments, production notes are IGNORED.
        if left:
            cell0    = cells[0]
            is_bold0 = any(r.bold for r in cell0.paragraphs[0].runs
                           if r.text.strip()) if cell0.paragraphs else False
            is_upper = left == left.upper() and len(left) > 2
            is_short = len(left.split()) <= 8
            if (is_bold0 or is_upper) and is_short and left != current_section:
                current_section = left
                lines.append({
                    "type": "section", "speaker": None,
                    "text": left, "notes": "",
                })

        # Right column
        if right:
            lines.extend(classify_content_cell(cells[1], ""))

    return lines



# ─────────────────────────────────────────────────────────────────────────────
# Audio extraction  (ffmpeg — available on Railway via Dockerfile)
# ─────────────────────────────────────────────────────────────────────────────

def extract_audio(video_path, out_wav):
    """Extract mono 16 kHz WAV from video (mixed down). Returns path."""
    cmd = [
        "ffmpeg", "-y", "-i", str(video_path),
        "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
        str(out_wav)
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg failed:\n{result.stderr[-800:]}")
    return out_wav


def extract_channel(video_path, out_mp3, channel=0):
    """
    Extract a single channel from a stereo file as MP3.
    channel=0 → left  (actuality/dialogue)
    channel=1 → right (VO narration)
    """
    pan = f"pan=mono|c0=c{channel}"
    cmd = [
        "ffmpeg", "-y", "-i", str(video_path),
        "-af", pan,
        "-ar", "16000", "-ac", "1",
        "-b:a", "64k",
        str(out_mp3)
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg channel extract failed:\n{result.stderr[-400:]}")
    return out_mp3


def is_stereo(video_path):
    """Return True if the audio has 2 or more channels."""
    cmd = [
        "ffprobe", "-v", "quiet",
        "-select_streams", "a:0",
        "-show_entries", "stream=channels",
        "-of", "csv=p=0",
        str(video_path)
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    try:
        return int(result.stdout.strip()) >= 2
    except:
        return False


def compress_audio(wav_path, out_mp3):
    """Compress WAV → MP3 at 64kbps to stay under OpenAI 25 MB limit."""
    cmd = [
        "ffmpeg", "-y", "-i", str(wav_path),
        "-b:a", "64k", str(out_mp3)
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg compress failed:\n{result.stderr[-400:]}")
    return out_mp3

# ─────────────────────────────────────────────────────────────────────────────
# EDL parser  (CMX 3600)
# ─────────────────────────────────────────────────────────────────────────────

def parse_edl(edl_path, fps=25):
    """
    Parse a CMX 3600 EDL file. Returns list of events in sequence order:
      { "event": int, "reel": str, "tc_in": str, "tc_out": str,
        "rec_in": str, "rec_out": str, "clip_name": str }

    tc_in/tc_out  = source timecode
    rec_in/rec_out = record (programme) timecode  ← this is what we want
    """
    events = []
    current = {}

    EDL_EVENT   = re.compile(
        r'^\s*(\d+)\s+(\S+)\s+\S+\s+\S+\s+'
        r'(\d{2}:\d{2}:\d{2}[:\;]\d{2})\s+'
        r'(\d{2}:\d{2}:\d{2}[:\;]\d{2})\s+'
        r'(\d{2}:\d{2}:\d{2}[:\;]\d{2})\s+'
        r'(\d{2}:\d{2}:\d{2}[:\;]\d{2})')
    CLIP_NAME   = re.compile(r'\*\s*(?:FROM CLIP NAME|CLIP NAME)\s*:\s*(.+)', re.IGNORECASE)
    LOC_NAME    = re.compile(r'\*\s*LOC\s*:\S+\s+\S+\s+(.+)', re.IGNORECASE)

    with open(edl_path, encoding="utf-8", errors="replace") as f:
        for line in f:
            line = line.rstrip()

            m = EDL_EVENT.match(line)
            if m:
                if current:
                    events.append(current)
                current = {
                    "event":     int(m.group(1)),
                    "reel":      m.group(2),
                    "src_in":    m.group(3).replace(";", ":"),
                    "src_out":   m.group(4).replace(";", ":"),
                    "rec_in":    m.group(5).replace(";", ":"),
                    "rec_out":   m.group(6).replace(";", ":"),
                    "clip_name": "",
                }
                continue

            # Clip name comment lines
            m = CLIP_NAME.match(line)
            if m and current:
                current["clip_name"] = m.group(1).strip()
                continue

            m = LOC_NAME.match(line)
            if m and current:
                if not current.get("clip_name"):
                    current["clip_name"] = m.group(1).strip()

    if current:
        events.append(current)

    # Remove duplicates (Avid sometimes outputs the same event twice)
    seen = set()
    unique = []
    for e in events:
        key = e["rec_in"]
        if key not in seen:
            seen.add(key)
            unique.append(e)

    return sorted(unique, key=lambda e: e["rec_in"])


def match_from_edl(script_lines, edl_events, fps):
    """
    Match VO lines to EDL events by sequence order.
    EDL event 1 → first VO line in script, event 2 → second VO line, etc.
    Actuality lines get blank TCs (not needed).
    """
    results   = []
    vo_lines  = [i for i, l in enumerate(script_lines) if l.get("type") == "vo"]
    edl_idx   = 0

    for i, line in enumerate(script_lines):
        ltype = line.get("type", "act")
        text  = line.get("text", "")

        if ltype in ("section", "part", "coda") or not text:
            results.append({**line, "tc_in": "", "tc_out": "", "dur": ""})
            continue

        if ltype != "vo":
            # Actuality — blank TC, reference only
            results.append({**line, "tc_in": "", "tc_out": "", "dur": ""})
            continue

        # VO line — assign next EDL event in sequence
        if edl_idx < len(edl_events):
            ev     = edl_events[edl_idx]
            tc_in  = ev["rec_in"]
            tc_out = ev["rec_out"]
            # Calculate duration in seconds
            def tc_secs(tc):
                p = tc.replace(";",":").split(":")
                return int(p[0])*3600 + int(p[1])*60 + int(p[2]) + int(p[3])/fps
            dur_secs = max(tc_secs(tc_out) - tc_secs(tc_in), 0)
            results.append({
                **line,
                "tc_in":  tc_in,
                "tc_out": tc_out,
                "dur":    dur_str(dur_secs),
            })
            edl_idx += 1
        else:
            # More VO lines than EDL events
            results.append({**line, "tc_in": "", "tc_out": "", "dur": ""})

    matched = sum(1 for r in results if r.get("tc_in") and r.get("type") == "vo")
    total_vo = sum(1 for l in script_lines if l.get("type") == "vo")
    print(f"EDL match: {matched} / {total_vo} VO lines assigned from {len(edl_events)} EDL events")
    return results




# ─────────────────────────────────────────────────────────────────────────────
# Whisper API transcription
# ─────────────────────────────────────────────────────────────────────────────

def transcribe_openai(audio_path, api_key):
    """
    Call OpenAI Whisper API with word-level timestamps.
    Returns list of {start, end, text} — one entry per segment,
    but start/end are taken from the first/last WORD in the segment
    for frame-accurate timing rather than segment-level estimates.
    """
    from openai import OpenAI

    client  = OpenAI(api_key=api_key)
    size_mb = Path(audio_path).stat().st_size / 1024 / 1024

    if size_mb > 24:
        raise RuntimeError(
            f"Audio file is {size_mb:.1f} MB — exceeds OpenAI 25 MB limit. "
            "Try a shorter programme or lower-bitrate export."
        )

    with open(audio_path, "rb") as f:
        response = client.audio.transcriptions.create(
            model                   = "whisper-1",
            file                    = f,
            response_format         = "verbose_json",
            timestamp_granularities = ["segment", "word"]
        )

    # Build word index: list of {start, end, word}
    words = []
    if hasattr(response, "words") and response.words:
        for w in response.words:
            words.append({
                "start": float(w.start),
                "end":   float(w.end),
                "word":  w.word.strip(),
            })

    # Map each segment to word-accurate start/end
    segments = []
    word_idx = 0
    for seg in response.segments:
        seg_start = float(seg.start)
        seg_end   = float(seg.end)
        seg_text  = seg.text.strip()

        if words:
            # Find words that fall within this segment
            seg_words = [w for w in words
                         if w["start"] >= seg_start - 0.1
                         and w["end"]   <= seg_end   + 0.1]
            if seg_words:
                seg_start = seg_words[0]["start"]
                seg_end   = seg_words[-1]["end"]

        segments.append({
            "start": seg_start,
            "end":   seg_end,
            "text":  seg_text,
        })

    return segments


# ─────────────────────────────────────────────────────────────────────────────
# Timecode matching
# ─────────────────────────────────────────────────────────────────────────────

def match_timecodes(script_lines, segments, tc_offset_secs, fps):
    """
    Dual-channel or mono matching:

    If `segments` is a dict {"dial": [...], "vo": [...]}, use:
      - dial segments to match actuality lines  (clean dialogue channel)
      - vo   segments to match VO lines         (clean narration channel)

    If `segments` is a plain list (mono/mixed), fall back to:
      - actuality-anchored matching with VO interpolation

    Confirmed TCs have no prefix. Uncertain/interpolated TCs get ~ prefix.
    """
    if not segments:
        return [{**l, "tc_in": "", "tc_out": "", "dur": ""} for l in script_lines]

    # ── Dual-channel path ─────────────────────────────────────────────────────
    if isinstance(segments, dict):
        dial_segs = segments.get("dial", [])
        vo_segs   = segments.get("vo",   [])

        WINDOW    = 12
        LOOKAHEAD = 120
        THRESHOLD = 0.20
        FALLBACK  = 0.08

        def best_match(norm_text, seg_list, cursor):
            n = len(seg_list)
            best_score = 0.0
            best_i = best_j = None
            end = min(cursor + LOOKAHEAD, n)
            for i in range(cursor, end):
                joined = ""
                for j in range(i, min(i + WINDOW, end)):
                    joined      = (joined + " " + seg_list[j]["text"]).strip()
                    norm_joined = normalize(joined)
                    score       = similarity(norm_text, norm_joined)
                    if len(norm_text) > 6 and norm_text in norm_joined:
                        score = max(score, 0.85)
                    if score > best_score:
                        best_score = score
                        best_i, best_j = i, j
            return best_score, best_i, best_j

        results       = []
        dial_cursor   = 0
        vo_cursor     = 0
        total_audio   = max(
            dial_segs[-1]["end"] if dial_segs else 0,
            vo_segs[-1]["end"]   if vo_segs   else 0
        )

        for line in script_lines:
            ltype = line.get("type", "act")
            text  = line.get("text", "").strip()

            if ltype in ("section", "part", "coda") or not text:
                results.append({**line, "tc_in": "", "tc_out": "", "dur": "",
                                "_t_in": None, "_matched": False})
                continue

            norm_text = normalize(text)

            # Choose which channel to match against
            if ltype == "vo":
                seg_list = vo_segs
                cursor   = vo_cursor
            else:
                seg_list = dial_segs
                cursor   = dial_cursor

            score, bi, bj = best_match(norm_text, seg_list, cursor)

            if score >= THRESHOLD and bi is not None:
                if ltype == "vo":
                    vo_cursor = bj + 1
                else:
                    dial_cursor = bj + 1
                t_in  = seg_list[bi]["start"]
                t_out = seg_list[bj]["end"]
                results.append({
                    **line,
                    "tc_in":  seconds_to_tc(t_in,  tc_offset_secs, fps),
                    "tc_out": seconds_to_tc(t_out, tc_offset_secs, fps),
                    "dur":    dur_str(t_out - t_in),
                    "_t_in":    t_in,
                    "_matched": True,
                })
            elif score >= FALLBACK and bi is not None:
                if ltype == "vo":
                    vo_cursor = bj + 1
                else:
                    dial_cursor = bj + 1
                t_in  = seg_list[bi]["start"]
                t_out = seg_list[bj]["end"]
                results.append({
                    **line,
                    "tc_in":  "~" + seconds_to_tc(t_in,  tc_offset_secs, fps),
                    "tc_out": "~" + seconds_to_tc(t_out, tc_offset_secs, fps),
                    "dur":    dur_str(t_out - t_in),
                    "_t_in":    t_in,
                    "_matched": True,
                })
            else:
                results.append({**line, "tc_in": "", "tc_out": "", "dur": "",
                                "_t_in": None, "_matched": False})

        # Interpolate any remaining blanks between anchors
        anchors = [(-1, 0.0)]
        for i, r in enumerate(results):
            if r.get("_t_in") is not None:
                anchors.append((i, r["_t_in"]))
        anchors.append((len(results), total_audio))

        for idx, r in enumerate(results):
            if r.get("_t_in") is not None:
                continue
            if r.get("type") in ("section","part","coda") or not r.get("text"):
                continue
            prev_a = anchors[0]
            next_a = anchors[-1]
            for a in anchors:
                if a[0] < idx:   prev_a = a
                elif a[0] > idx: next_a = a; break
            prev_i, prev_t = prev_a
            next_i, next_t = next_a
            span = max(next_i - prev_i, 1)
            frac = (idx - prev_i) / span
            est_in  = prev_t + frac * (next_t - prev_t)
            est_out = min(est_in + 4.0, next_t)
            results[idx]["tc_in"]  = "~" + seconds_to_tc(est_in,  tc_offset_secs, fps)
            results[idx]["tc_out"] = "~" + seconds_to_tc(est_out, tc_offset_secs, fps)
            results[idx]["dur"]    = dur_str(est_out - est_in)
            results[idx]["_t_in"]  = est_in

        n_conf  = sum(1 for r in results if r.get("_matched"))
        n_interp = sum(1 for r in results if not r.get("_matched") and r.get("tc_in"))
        for r in results:
            r.pop("_matched", None); r.pop("_t_in", None)
        print(f"Dual-channel match: {n_conf} confirmed | {n_interp} interpolated")
        return results

    # ── Mono / mixed fallback path ────────────────────────────────────────────
    n_segs     = len(segments)
    seg_cursor = 0
    WINDOW     = 12
    LOOKAHEAD  = 120
    THRESHOLD  = 0.20
    FALLBACK   = 0.10
    total_audio = segments[-1]["end"]

    # ── Pass 1: match actuality lines only ───────────────────────────────────
    results = []
    for line in script_lines:
        ltype = line.get("type", "act")
        text  = line.get("text", "").strip()

        # Section headers — no TC
        if ltype in ("section", "part", "coda") or not text:
            results.append({**line, "tc_in": "", "tc_out": "", "dur": "",
                            "_t_in": None, "_t_out": None, "_matched": False})
            continue

        # VO lines — skip in pass 1, will be interpolated in pass 2
        if ltype == "vo":
            results.append({**line, "tc_in": "", "tc_out": "", "dur": "",
                            "_t_in": None, "_t_out": None, "_matched": False})
            continue

        # Actuality — try to match
        norm_line  = normalize(text)
        best_score = 0.0
        best_i = best_j = None
        search_end = min(seg_cursor + LOOKAHEAD, n_segs)

        for i in range(seg_cursor, search_end):
            joined = ""
            for j in range(i, min(i + WINDOW, search_end)):
                joined      = (joined + " " + segments[j]["text"]).strip()
                norm_joined = normalize(joined)
                score       = similarity(norm_line, norm_joined)
                if len(norm_line) > 6 and norm_line in norm_joined:
                    score = max(score, 0.85)
                if score > best_score:
                    best_score = score
                    best_i, best_j = i, j

        if best_score >= THRESHOLD and best_i is not None:
            seg_cursor = best_j + 1
            t_in  = segments[best_i]["start"]
            t_out = segments[best_j]["end"]
            results.append({
                **line,
                "tc_in":  seconds_to_tc(t_in,  tc_offset_secs, fps),
                "tc_out": seconds_to_tc(t_out, tc_offset_secs, fps),
                "dur":    dur_str(t_out - t_in),
                "_t_in":    t_in,
                "_t_out":   t_out,
                "_matched": True,
            })
        elif best_score >= FALLBACK and best_i is not None:
            seg_cursor = best_j + 1
            t_in  = segments[best_i]["start"]
            t_out = segments[best_j]["end"]
            results.append({
                **line,
                "tc_in":  "~" + seconds_to_tc(t_in,  tc_offset_secs, fps),
                "tc_out": "~" + seconds_to_tc(t_out, tc_offset_secs, fps),
                "dur":    dur_str(t_out - t_in),
                "_t_in":    t_in,
                "_t_out":   t_out,
                "_matched": True,
            })
        else:
            results.append({**line, "tc_in": "", "tc_out": "", "dur": "",
                            "_t_in": None, "_t_out": None, "_matched": False})

    # ── Pass 2: place VO lines between actuality anchors ─────────────────────
    # Build list of (index, t_in) for every line that has a real timestamp
    anchors = [(-1, 0.0)]
    for i, r in enumerate(results):
        if r.get("_t_in") is not None:
            anchors.append((i, r["_t_in"]))
    anchors.append((len(results), total_audio))

    for idx, r in enumerate(results):
        ltype = r.get("type", "act")
        if r.get("_t_in") is not None:
            continue   # already has a TC
        if ltype in ("section", "part", "coda") or not r.get("text"):
            continue   # no TC for headers

        # Find surrounding anchors
        prev_a = anchors[0]
        next_a = anchors[-1]
        for a in anchors:
            if a[0] < idx:
                prev_a = a
            elif a[0] > idx:
                next_a = a
                break

        prev_i, prev_t = prev_a
        next_i, next_t = next_a
        span_lines = max(next_i - prev_i, 1)
        span_time  = max(next_t - prev_t, 0.0)
        frac       = (idx - prev_i) / span_lines
        est_in     = prev_t + frac * span_time

        # Estimate duration: typical VO line ~3-8 seconds
        # Use available gap divided by number of VO lines in this block
        vo_in_block = sum(
            1 for r2 in results[prev_i+1:next_i]
            if r2.get("type") == "vo" and r2.get("_t_in") is None
        )
        est_dur = min(span_time / max(vo_in_block, 1), 8.0) if span_time > 0 else 4.0
        est_out = min(est_in + est_dur, next_t)

        results[idx]["tc_in"]  = "~" + seconds_to_tc(est_in,  tc_offset_secs, fps)
        results[idx]["tc_out"] = "~" + seconds_to_tc(est_out, tc_offset_secs, fps)
        results[idx]["dur"]    = dur_str(est_out - est_in)
        results[idx]["_t_in"]  = est_in

    # Clean internal keys
    n_act_matched = sum(1 for r in results
                        if r.get("_matched") and r["type"] == "act")
    n_vo_interp   = sum(1 for r in results
                        if r.get("type") == "vo" and r.get("tc_in","").startswith("~"))
    n_vo_blank    = sum(1 for r in results
                        if r.get("type") == "vo" and not r.get("tc_in"))
    for r in results:
        r.pop("_matched", None); r.pop("_t_in", None); r.pop("_t_out", None)

    print(f"Match: {n_act_matched} actualities matched | "
          f"{n_vo_interp} VO interpolated | {n_vo_blank} blank")
    return results



# ─────────────────────────────────────────────────────────────────────────────
# Output .docx builder
# ─────────────────────────────────────────────────────────────────────────────


def match_three_input(script_lines, edl_events, audio_segments,
                      tc_offset_secs, fps, log_fn=print):
    """
    Matching strategy:

    PRIMARY  — Whisper audio (when available).
               Find each VO line in the transcript by text similarity.
               Uses word-level timestamps for accuracy.

    VERIFY   — EDL cross-check (when available).
               After Whisper finds a position, check if an EDL event
               falls within 2 seconds. If yes, use the EDL TC (frame-
               accurate). If no nearby EDL event, keep Whisper TC (~).

    FALLBACK — Sequential EDL (no audio at all).
               EDL event 1 → VO line 1, event 2 → VO line 2, etc.

    Actuality lines always get blank TCs.
    Unmatched VO lines get interpolated TCs (~) between confirmed anchors.
    """
    WORDING_THRESHOLD = 0.15
    EDL_VERIFY_SECS   = 2.0
    LOOKAHEAD         = 80
    WINDOW            = 10
    THRESHOLD         = 0.20
    FALLBACK_SCORE    = 0.10

    results = []

    # ── Flatten audio ──────────────────────────────────────────────────────
    if isinstance(audio_segments, dict):
        vo_segs = audio_segments.get("vo", [])
    elif audio_segments:
        vo_segs = audio_segments
    else:
        vo_segs = []

    n_vo      = len(vo_segs)
    vo_cursor = 0
    total_dur = vo_segs[-1]["end"] if vo_segs else 0.0

    # ── EDL lookup by time ────────────────────────────────────────────────
    def tc_s(tc):
        p = tc.replace(";", ":").split(":")
        return int(p[0])*3600 + int(p[1])*60 + int(p[2]) + int(p[3])/fps

    edl_by_time = []
    for ev in edl_events:
        try:
            edl_by_time.append((tc_s(ev["rec_in"]), tc_s(ev["rec_out"]), ev))
        except:
            pass
    edl_by_time.sort(key=lambda x: x[0])
    edl_used = set()

    def find_edl_near(w_in):
        """Find unused EDL event within EDL_VERIFY_SECS of Whisper TC In."""
        best = None; best_diff = EDL_VERIFY_SECS + 1
        for t_in, t_out, ev in edl_by_time:
            if id(ev) in edl_used: continue
            diff = abs(t_in - w_in)
            if diff <= EDL_VERIFY_SECS and diff < best_diff:
                best_diff = diff; best = (t_in, t_out, ev)
        return best

    def best_whisper(norm_text):
        nonlocal vo_cursor
        if vo_cursor >= n_vo:
            return 0.0, None, None
        best_score = 0.0; best_i = best_j = None
        end = min(vo_cursor + LOOKAHEAD, n_vo)
        for i in range(vo_cursor, end):
            joined = ""
            for j in range(i, min(i + WINDOW, end)):
                joined = (joined + " " + vo_segs[j]["text"]).strip()
                score  = similarity(normalize(norm_text), normalize(joined))
                if len(norm_text) > 6 and normalize(norm_text) in normalize(joined):
                    score = max(score, 0.85)
                if score > best_score:
                    best_score = score; best_i = i; best_j = j
        return best_score, best_i, best_j

    # ── Main pass ─────────────────────────────────────────────────────────
    edl_seq_idx = 0   # for sequential EDL fallback (no audio)

    for line in script_lines:
        ltype = line.get("type", "act")
        text  = line.get("text", "").strip()

        if ltype in ("section", "part", "coda") or not text:
            results.append({**line, "tc_in": "", "tc_out": "", "dur": "",
                            "_match_note": "", "_t_in": None})
            continue

        if ltype != "vo":
            results.append({**line, "tc_in": "", "tc_out": "", "dur": "",
                            "_match_note": "", "_t_in": None})
            continue

        tc_in = tc_out = dur = note = ""
        t_in_raw = None

        if vo_segs:
            # ── PRIMARY: Whisper match ─────────────────────────────────
            score, bi, bj = best_whisper(normalize(text))

            if score >= THRESHOLD and bi is not None:
                vo_cursor = bj + 1
                w_in  = vo_segs[bi]["start"]
                w_out = vo_segs[bj]["end"]

                # Wording check
                heard = " ".join(vo_segs[k]["text"].strip()
                                 for k in range(bi, bj+1)).strip()
                sim = similarity(normalize(text), normalize(heard))
                if sim < (1 - WORDING_THRESHOLD):
                    s = text[:60] + ("..." if len(text) > 60 else "")
                    h = heard[:60] + ("..." if len(heard) > 60 else "")
                    note = f"WORDING: script: {s!r} | audio: {h!r}"

                # VERIFY: check EDL
                edl_hit = find_edl_near(w_in)
                if edl_hit:
                    t_in_edl, t_out_edl, ev = edl_hit
                    edl_used.add(id(ev))
                    tc_in    = ev["rec_in"]
                    tc_out   = ev["rec_out"]
                    dur      = dur_str(max(t_out_edl - t_in_edl, 0))
                    t_in_raw = t_in_edl
                else:
                    tc_in    = "~" + seconds_to_tc(w_in,  tc_offset_secs, fps)
                    tc_out   = "~" + seconds_to_tc(w_out, tc_offset_secs, fps)
                    dur      = dur_str(w_out - w_in)
                    t_in_raw = w_in

            elif score >= FALLBACK_SCORE and bi is not None:
                vo_cursor = bj + 1
                w_in  = vo_segs[bi]["start"]
                w_out = vo_segs[bj]["end"]
                tc_in    = "~" + seconds_to_tc(w_in,  tc_offset_secs, fps)
                tc_out   = "~" + seconds_to_tc(w_out, tc_offset_secs, fps)
                dur      = dur_str(w_out - w_in)
                t_in_raw = w_in
                note     = "⚠ low-confidence match"
            else:
                note = "NOT FOUND in audio"

        elif edl_events:
            # ── FALLBACK: sequential EDL (no audio) ───────────────────
            if edl_seq_idx < len(edl_events):
                ev       = edl_events[edl_seq_idx]; edl_seq_idx += 1
                t_in_e   = tc_s(ev["rec_in"]); t_out_e = tc_s(ev["rec_out"])
                tc_in    = ev["rec_in"]; tc_out = ev["rec_out"]
                dur      = dur_str(max(t_out_e - t_in_e, 0))
                t_in_raw = t_in_e
            else:
                note = "NOT FOUND — EDL exhausted"
        else:
            note = "NOT FOUND — no audio or EDL"

        results.append({**line, "tc_in": tc_in, "tc_out": tc_out, "dur": dur,
                        "_match_note": note, "_t_in": t_in_raw})

    # ── Interpolation ─────────────────────────────────────────────────────
    anchors = [(-1, 0.0)]
    for i, r in enumerate(results):
        if r.get("_t_in") is not None:
            anchors.append((i, r["_t_in"]))
    anchors.append((len(results), total_dur or (
        tc_s(edl_events[-1]["rec_out"]) if edl_events else 3600.0)))

    for idx, r in enumerate(results):
        if r.get("_t_in") is not None: continue
        if r.get("type") in ("section","part","coda"): continue
        if r.get("type") != "vo": continue
        if "NOT FOUND" in r.get("_match_note", ""): continue

        prev_a = anchors[0]; next_a = anchors[-1]
        for a in anchors:
            if a[0] < idx:   prev_a = a
            elif a[0] > idx: next_a = a; break
        prev_i, prev_t = prev_a; next_i, next_t = next_a
        frac    = (idx - prev_i) / max(next_i - prev_i, 1)
        est_in  = prev_t + frac * (next_t - prev_t)
        est_out = min(est_in + 5.0, next_t)
        results[idx].update({
            "tc_in":       "~" + seconds_to_tc(est_in,  tc_offset_secs, fps),
            "tc_out":      "~" + seconds_to_tc(est_out, tc_offset_secs, fps),
            "dur":         dur_str(est_out - est_in),
            "_match_note": "⚠ TC estimated — check against cut",
            "_t_in":       est_in,
        })

    n_edl    = sum(1 for r in results if r.get("type")=="vo"
                   and r.get("tc_in") and not r["tc_in"].startswith("~")
                   and "NOT FOUND" not in r.get("_match_note",""))
    n_whisper = sum(1 for r in results if r.get("type")=="vo"
                    and r.get("tc_in","").startswith("~")
                    and "NOT FOUND" not in r.get("_match_note","")
                    and "estimated" not in r.get("_match_note",""))
    n_interp = sum(1 for r in results if "estimated" in r.get("_match_note",""))
    n_warn   = sum(1 for r in results if "WORDING" in r.get("_match_note",""))
    n_miss   = sum(1 for r in results if "NOT FOUND" in r.get("_match_note",""))
    log_fn(f"Match: {n_edl} EDL-verified | {n_whisper} Whisper-only | "
           f"{n_interp} estimated | {n_warn} wording ⚠ | {n_miss} missing")

    for r in results:
        r.pop("_t_in", None)
    return results


def build_output_xlsx(matched_lines, output_path, fps):
    """
    Build the VO script as an Excel workbook.
    Columns: TC In | TC Out | Duration | Script & VO | Notes
    Colour coding:
      #B8CCE4 — story/segment header (blue)
      #BFBFBF — part/act break (mid grey)
      #D9D9D9 — actuality (light grey)
      white   — VO
    Calibri 10pt throughout. Notes column blank for manual entry.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "VO Script"

    # ── Style helpers ────────────────────────────────────────────────────────
    def font(bold=False, italic=False, color="000000"):
        return Font(name="Calibri", size=10, bold=bold,
                    italic=italic, color=color)

    def fill(hex_color):
        return PatternFill("solid", fgColor=hex_color.replace("#",""))

    def border():
        s = Side(style="thin", color="AAAAAA")
        return Border(left=s, right=s, top=s, bottom=s)

    def align(wrap=True, horizontal="left"):
        return Alignment(wrap_text=wrap, vertical="top",
                         horizontal=horizontal)

    FILLS = {
        "section": fill("B8CCE4"),
        "part":    fill("BFBFBF"),
        "act":     fill("D9D9D9"),
        "vo":      fill("FFFFFF"),
        "header":  fill("1F3864"),
        "coda":    fill("E2EFD9"),
    }

    # ── Column widths (chars) ────────────────────────────────────────────────
    ws.column_dimensions["A"].width = 15   # TC In
    ws.column_dimensions["B"].width = 15   # TC Out
    ws.column_dimensions["C"].width = 14   # Duration
    ws.column_dimensions["D"].width = 60   # Script & VO
    ws.column_dimensions["E"].width = 30   # Notes

    # ── Header row ────────────────────────────────────────────────────────────
    headers = ["TC IN", "TC OUT", "DURATION", "SCRIPT & VO", "NOTES"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font      = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
        cell.fill      = FILLS["header"]
        cell.border    = border()
        cell.alignment = align(horizontal="center")
    ws.row_dimensions[1].height = 18
    ws.freeze_panes = "A2"

    # ── Data rows ─────────────────────────────────────────────────────────────
    row_num   = 2
    prev_type = None   # for actuality grouping

    for line in matched_lines:
        ltype      = line.get("type",        "act")
        text       = line.get("text",        "").strip()
        tc_in      = line.get("tc_in",       "")
        tc_out     = line.get("tc_out",      "")
        dur        = line.get("dur",         "")
        speaker    = line.get("speaker",     "")
        match_note = line.get("_match_note", "")
        notes      = match_note   # col 1 source notes are suppressed

        # ── Skip COMM rows and bare "VO" rows ─────────────────────────────
        if text.upper() == "COMM":
            continue
        if ltype == "vo" and text.upper() in ("", "VO"):
            continue
        if not text:
            continue

        # ── Section / part / coda — blue header, no merged cells ──────────
        if ltype in ("section", "part", "coda"):
            label_text = text.upper() if ltype == "section" else text
            for col in range(1, 6):
                cell = ws.cell(row=row_num, column=col,
                               value=label_text if col == 4 else "")
                cell.fill      = FILLS["section"]
                cell.border    = border()
                cell.alignment = align(horizontal="left")
                cell.font      = Font(name="Calibri", size=10,
                                      bold=(ltype != "coda"),
                                      italic=(ltype == "coda"),
                                      color="000000")
            ws.row_dimensions[row_num].height = 16
            prev_type = ltype
            row_num += 1
            continue

        # ── VO row ────────────────────────────────────────────────────────
        if ltype == "vo":
            display_in  = tc_in.lstrip("~")
            display_out = tc_out.lstrip("~")
            if tc_in.startswith("~") and not notes:
                notes = "⚠ TC estimated — check against cut"
            try:
                def _s(tc):
                    p = tc.lstrip("~").replace(";",":").split(":")
                    return int(p[0])*3600+int(p[1])*60+int(p[2])+int(p[3])/25
                if dur and (_s(display_out) - _s(display_in)) < 1.0:
                    if not notes: notes = "⚠ Duration <1s — check TC"
            except: pass

            script_text = "VO: " + text

            values = [display_in, display_out, dur, script_text, notes]
            for col, val in enumerate(values, 1):
                cell = ws.cell(row=row_num, column=col, value=val)
                cell.fill      = FILLS["vo"]
                cell.border    = border()
                cell.alignment = align()
                if col == 4:
                    cell.font = Font(name="Calibri", size=10, bold=True)
                elif col in (1, 2, 3):
                    cell.font = Font(name="Courier New", size=10)
                else:
                    is_warn   = bool(notes and "WORDING" in notes)
                    cell.font = Font(name="Calibri", size=10,
                                     italic=is_warn,
                                     color="CC0000" if is_warn else "000000")
            if notes and "WORDING" in notes:
                ws.cell(row=row_num, column=5).font = Font(
                    name="Calibri", size=10, color="CC0000")
            ws.row_dimensions[row_num].height = 15
            prev_type = ltype
            row_num += 1
            continue

        # ── Actuality row ─────────────────────────────────────────────────
        # First line of a consecutive actuality group gets "Actuality:" prefix
        is_first_act = (prev_type != "act")
        if speaker:
            if is_first_act:
                script_text = f"Actuality: {speaker}: {text}"
            else:
                script_text = f"{speaker}: {text}"
        else:
            if is_first_act:
                script_text = f"Actuality: {text}"
            else:
                script_text = text

        for col, val in enumerate(["", "", "", script_text, ""], 1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.fill      = FILLS["act"]
            cell.border    = border()
            cell.alignment = align()
            cell.font      = Font(name="Calibri", size=10,
                                  italic=(col == 4), color="444444" if col == 4 else "000000")

        ws.row_dimensions[row_num].height = 15
        prev_type = ltype
        row_num += 1

    wb.save(str(output_path))



# ─────────────────────────────────────────────────────────────────────────────
# Email notifications via Resend
# ─────────────────────────────────────────────────────────────────────────────

def send_notify_email(user_name, script_name, video_name,
                      n_matched, n_total, status, error):
    """
    Send a job completion/failure email via Resend API.
    Requires env vars:
      RESEND_API_KEY  — from resend.com (free tier: 3000 emails/month)
      NOTIFY_EMAIL    — address to send notifications TO
      NOTIFY_FROM     — sending address (must be verified in Resend)
                        e.g. "Stampede Formatter <notifications@yourdomain.com>"
                        or use Resend's default: "onboarding@resend.dev" for testing
    """
    import urllib.request, urllib.error, json as _json

    api_key    = os.environ.get("RESEND_API_KEY", "").strip()
    to_email   = os.environ.get("NOTIFY_EMAIL",   "").strip()
    from_addr  = os.environ.get("NOTIFY_FROM",    "Stampede VO Formatter <onboarding@resend.dev>").strip()

    if not api_key or not to_email:
        return  # Silently skip if not configured

    from datetime import datetime, timezone
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    if status == "done":
        subject = f"✓ VO Script ready — {user_name} — {script_name}"
        body = f"""<html><body style="font-family:Arial,sans-serif;color:#222;padding:24px">
<h2 style="color:#1F3864">Stampede VO Script Auto Formatter</h2>
<p style="font-size:15px">A new script has been formatted successfully.</p>
<table style="border-collapse:collapse;margin:20px 0;font-size:14px">
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">User</td>
      <td style="padding:6px 0">{user_name}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Script</td>
      <td style="padding:6px 0">{script_name}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Video</td>
      <td style="padding:6px 0">{video_name}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Lines matched</td>
      <td style="padding:6px 0">{n_matched} / {n_total}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Time</td>
      <td style="padding:6px 0">{now}</td></tr>
</table>
<p style="color:#888;font-size:12px">Stampede Productions · VO Script Auto Formatter</p>
</body></html>"""
    else:
        subject = f"✗ VO Formatter error — {user_name} — {script_name}"
        body = f"""<html><body style="font-family:Arial,sans-serif;color:#222;padding:24px">
<h2 style="color:#c0392b">Stampede VO Formatter — Job Failed</h2>
<table style="border-collapse:collapse;margin:20px 0;font-size:14px">
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">User</td>
      <td>{user_name}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Script</td>
      <td>{script_name}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Video</td>
      <td>{video_name}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Error</td>
      <td style="color:#c0392b">{error}</td></tr>
  <tr><td style="padding:6px 16px 6px 0;color:#666;font-weight:bold">Time</td>
      <td>{now}</td></tr>
</table>
</body></html>"""

    payload = _json.dumps({
        "from":    from_addr,
        "to":      [to_email],
        "subject": subject,
        "html":    body,
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.resend.com/emails",
        data    = payload,
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type":  "application/json",
        },
        method = "POST"
    )
    try:
        urllib.request.urlopen(req, timeout=10)
    except Exception as e:
        print(f"Email send failed (non-fatal): {e}")


# ─────────────────────────────────────────────────────────────────────────────
# Background job runner
# ─────────────────────────────────────────────────────────────────────────────

def run_job(job_id, script_path, video_path, tc_offset, fps, api_key, output_path, edl_path=None):
    job = jobs[job_id]

    def log(msg, pct=None):
        job["log"].append(msg)
        if pct is not None:
            job["progress"] = pct
        print(f"[{job_id[:8]}] {msg}")

    tmp_dir = Path(tempfile.mkdtemp())
    try:
        # 1 — Parse script
        log("Parsing source script…", 5)
        script_lines = parse_source_script(script_path)
        log(f"✓ Found {len(script_lines)} lines in script.", 15)

        # 2 — Parse EDL if provided
        edl_events = []
        if edl_path:
            log("Parsing EDL…", 18)
            edl_events = parse_edl(edl_path, fps)
            log(f"✓ {len(edl_events)} EDL events parsed", 22)

        # 3 — Transcribe audio if provided (fallback + wording check)
        audio_segments = None

        # 2 — Audio: extract from video, or use directly if already audio
        AUDIO_EXTS = {".mp3", ".mp4a", ".m4a", ".aac", ".wav", ".flac", ".ogg", ".opus"}
        upload_ext = Path(video_path).suffix.lower()

        mp3      = tmp_dir / "audio.mp3"
        mp3_dial = tmp_dir / "audio_dial.mp3"
        mp3_vo   = tmp_dir / "audio_vo.mp3"

        if upload_ext in AUDIO_EXTS and upload_ext == ".mp3":
            size_mb = Path(video_path).stat().st_size / 1024 / 1024
            log(f"Audio file detected ({size_mb:.1f} MB)…", 20)
            if size_mb > 24:
                log("File over 25 MB — re-compressing…", 25)
                compress_audio(video_path, mp3)
            else:
                import shutil as _sh
                _sh.copy2(video_path, mp3)
        elif upload_ext in AUDIO_EXTS:
            size_mb = Path(video_path).stat().st_size / 1024 / 1024
            log(f"Audio file detected ({upload_ext}, {size_mb:.1f} MB) — converting…", 20)
            compress_audio(video_path, mp3)
        else:
            log("Extracting audio from video (ffmpeg)…", 20)
            wav = tmp_dir / "audio.wav"
            extract_audio(video_path, wav)
            log("Compressing audio…", 28)
            compress_audio(wav, mp3)

        size_mb = mp3.stat().st_size / 1024 / 1024
        log(f"✓ Audio ready — {size_mb:.1f} MB", 30)

        # 3 — Transcribe audio if we have it
        if video_path:
            stereo = is_stereo(video_path)
            if stereo:
                log("Stereo detected — extracting L (dialogue) and R (VO) channels…", 32)
                extract_channel(video_path, mp3_dial, channel=0)
                extract_channel(video_path, mp3_vo,   channel=1)
                log(f"Dialogue: {mp3_dial.stat().st_size/1024/1024:.1f} MB  VO: {mp3_vo.stat().st_size/1024/1024:.1f} MB", 36)
                log("Transcribing dialogue channel…", 38)
                dial_segs = transcribe_openai(mp3_dial, api_key)
                log(f"✓ Dialogue: {len(dial_segs)} segments", 52)
                log("Transcribing VO channel…", 54)
                vo_segs = transcribe_openai(mp3_vo, api_key)
                log(f"✓ VO: {len(vo_segs)} segments", 68)
                audio_segments = {"dial": dial_segs, "vo": vo_segs}
            else:
                log("Mono/mixed audio — transcribing…", 36)
                audio_segments = transcribe_openai(mp3, api_key)
                log(f"✓ {len(audio_segments)} segments", 68)
        else:
            log("No audio file — using EDL only", 68)

        # 4 — Match: EDL first, audio fallback, wording diff check
        log("Matching timecodes to script…", 75)
        offset = tc_to_seconds(tc_offset)
        matched = match_three_input(
            script_lines, edl_events, audio_segments, offset, fps,
            log_fn=lambda m: log(m)
        )
        n_matched = sum(1 for m in matched
                        if m.get("tc_in") and m.get("type") == "vo")
        log(f"✓ {n_matched} VO lines timecoded", 88)

        # 5 — Build xlsx
        log("Building output spreadsheet…", 92)
        build_output_xlsx(matched, output_path, fps)
        log("✓ Done!", 100)

        job["status"]      = "done"
        job["output_path"] = str(output_path)

        # Send completion email
        send_notify_email(
            user_name    = job.get("user_name", "unknown"),
            script_name  = job.get("script_name", ""),
            video_name   = job.get("video_name", ""),
            n_matched    = n_matched,
            n_total      = len(matched),
            status       = "done",
            error        = None,
        )

    except Exception as e:
        job["status"] = "error"
        job["error"]  = str(e)
        log(f"✗ Error: {e}")
        send_notify_email(
            user_name   = job.get("user_name", "unknown"),
            script_name = job.get("script_name", ""),
            video_name  = job.get("video_name", ""),
            n_matched   = 0,
            n_total     = 0,
            status      = "error",
            error       = str(e),
        )
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        # Clean up uploaded files
        try: os.remove(script_path)
        except: pass
        try: os.remove(video_path)
        except: pass


# ─────────────────────────────────────────────────────────────────────────────
# Flask routes
# ─────────────────────────────────────────────────────────────────────────────

UPLOAD_DIR = Path(tempfile.gettempdir()) / "pp_timecoder_uploads"
OUTPUT_DIR = Path(tempfile.gettempdir()) / "pp_timecoder_outputs"
LOG_FILE   = Path(tempfile.gettempdir()) / "pp_timecoder_usage.csv"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

PERMANENT_API_KEY = os.environ.get("OPENAI_API_KEY", "").strip()


def write_usage_log(user_name, script_filename, video_filename, status, note=""):
    is_new = not LOG_FILE.exists()
    try:
        with open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            if is_new:
                w.writerow(["timestamp_utc", "user_name", "script_file", "video_file", "status", "note"])
            w.writerow([
                datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S"),
                user_name or "unknown",
                script_filename, video_filename, status, note,
            ])
    except Exception as e:
        print(f"Usage log write failed: {e}")


@app.route("/")
def index():
    return render_template("index.html", has_permanent_key=bool(PERMANENT_API_KEY))


@app.route("/api/start", methods=["POST"])
def start_job():
    if "script" not in request.files:
        return jsonify({"error": "Source script file is required."}), 400
    # Either a video/audio file OR an EDL must be provided
    has_audio = "video" in request.files and request.files["video"].filename
    has_edl   = "edl"   in request.files and request.files["edl"].filename
    if not has_audio and not has_edl:
        return jsonify({"error": "Please upload either an audio/video file or an EDL file."}), 400

    script_file = request.files["script"]
    video_file  = request.files.get("video")
    edl_file    = request.files.get("edl")
    user_name   = request.form.get("user_name", "").strip()
    tc_offset   = request.form.get("tc_offset", "10:00:00:00").strip()
    fps         = float(request.form.get("fps", "25"))

    if PERMANENT_API_KEY:
        api_key = PERMANENT_API_KEY
    else:
        api_key = request.form.get("api_key", "").strip()
        if not api_key:
            return jsonify({"error": "OpenAI API key is required."}), 400
        if not api_key.startswith("sk-"):
            return jsonify({"error": "That does not look like a valid OpenAI API key."}), 400

    if not user_name:
        return jsonify({"error": "Please enter your name so we can track usage."}), 400

    job_id      = str(uuid.uuid4())
    script_path = UPLOAD_DIR / f"{job_id}_script.docx"
    video_path  = UPLOAD_DIR / f"{job_id}_video{Path(video_file.filename).suffix}" if (video_file and video_file.filename) else None
    output_path = OUTPUT_DIR / f"{job_id}_timecoded.xlsx"

    script_file.save(str(script_path))
    if video_file and video_file.filename:
        video_file.save(str(video_path))
    edl_path = None
    if edl_file and edl_file.filename:
        edl_path = str(UPLOAD_DIR / f"{job_id}_edl{Path(edl_file.filename).suffix}")
        edl_file.save(edl_path)
    write_usage_log(user_name, script_file.filename,
                    (video_file.filename if video_file else edl_file.filename if edl_file else ""), "started")

    jobs[job_id] = {
        "status":      "running",
        "progress":    0,
        "log":         [],
        "output_path": None,
        "error":       None,
        "user_name":   user_name,
        "script_name": script_file.filename,
        "video_name":  video_file.filename,
    }

    thread = threading.Thread(
        target=run_job,
        args=(job_id, str(script_path),
              str(video_path) if (video_file and video_file.filename) else None,
              tc_offset, fps, api_key, output_path, edl_path),
        daemon=True
    )
    thread.start()
    return jsonify({"job_id": job_id})


@app.route("/api/status/<job_id>")
def job_status(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    if job["status"] in ("done", "error") and not job.get("_logged"):
        job["_logged"] = True
        write_usage_log(
            job.get("user_name"), job.get("script_name"), job.get("video_name"),
            job["status"],
            job.get("error", "")[:200] if job["status"] == "error" else ""
        )
    return jsonify({"status": job["status"], "progress": job["progress"],
                    "log": job["log"], "error": job.get("error")})


@app.route("/api/download/<job_id>")
def download(job_id):
    job = jobs.get(job_id)
    if not job or job["status"] != "done":
        return jsonify({"error": "Not ready"}), 404
    path = job["output_path"]
    if not path or not Path(path).exists():
        return jsonify({"error": "Output file missing"}), 500
    return send_file(path, as_attachment=True, download_name="timecoded_script.xlsx")


@app.route("/usage")
def usage_log():
    pw = os.environ.get("USAGE_PASSWORD", "")
    if pw and request.args.get("key") != pw:
        return "Unauthorised", 403
    if not LOG_FILE.exists():
        return "<p>No usage data yet.</p>"
    with open(LOG_FILE, encoding="utf-8") as f:
        rows = list(csv.reader(f))
    if len(rows) < 2:
        return "<p>No usage data yet.</p>"
    headers, data = rows[0], list(reversed(rows[1:]))
    tbl = "".join("<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>" for row in data)
    hdr = "".join(f"<th>{h}</th>" for h in headers)
    return f"""<!DOCTYPE html><html><head><title>Usage Log</title>
<style>body{{font-family:Arial,sans-serif;padding:30px;background:#0f0f1a;color:#e8e8f0}}
h1{{color:#5b6ef5;margin-bottom:20px}}table{{border-collapse:collapse;width:100%;font-size:13px}}
th{{background:#1f3864;color:#fff;padding:10px 14px;text-align:left}}
td{{padding:8px 14px;border-bottom:1px solid #2e2e50}}tr:hover td{{background:#1a1a2e}}</style>
</head><body><h1>Stampede VO Formatter — Usage Log</h1>
<p style=color:#8080a8>{len(data)} jobs total</p>
<table><thead><tr>{hdr}</tr></thead><tbody>{tbl}</tbody></table></body></html>"""


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
